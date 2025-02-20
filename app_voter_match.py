# App candidate match

"""
    Note: The app allows users to select their stance on various topics and then compares their stances with those of the candidates.
        - The app uses a pre-trained NLI model to compare user stances with candidate statements.
        - The app then aggregates the scores for each candidate to provide an overall alignment score.
        - The user can see the alignment for each topic and the overall alignment with each candidate.

    Run : python app_voter_match.py

"""

# Import necessary libraries
import re
from docx import Document
from collections import defaultdict
from nltk.corpus import stopwords
from nltk.tokenize import TreebankWordTokenizer
from nltk.stem import WordNetLemmatizer
import nltk
import os
import json
from docx import Document
import gradio as gr
import logging


# # Ensure NLTK resources are downloaded
#nltk.download('stopwords')
#nltk.download('wordnet')
#nltk.download('punkt')

nltk_data_dir = os.path.join(os.path.expanduser('~'), 'nltk_data')

if not os.path.exists(nltk_data_dir):
    os.makedirs(nltk_data_dir)

nltk.data.path.append(nltk_data_dir)

try:
    stop_words = set(stopwords.words('english'))
except LookupError:
    nltk.download('stopwords', download_dir=nltk_data_dir)
    stop_words = set(stopwords.words('english'))

try:
    nltk.data.find('corpora/wordnet.zip')
except LookupError:
    nltk.download('wordnet', download_dir=nltk_data_dir)

try:
    nltk.data.find('tokenizers/punkt.zip')
except LookupError:
    nltk.download('punkt', download_dir=nltk_data_dir)

try:
    nltk.data.find('corpora/omw-1.4.zip')
except LookupError:
    nltk.download('omw-1.4', download_dir=nltk_data_dir)



# Initialize stopwords, lemmatizer, and tokenizer
#stop_words = set(stopwords.words('english'))
lemmatizer = WordNetLemmatizer()
tokenizer = TreebankWordTokenizer()

# 
def preprocess_text(text):
    # Lowercasing
    text = text.lower()
    
    # Text Standardization
    replacements = {
        "black lives matter": "black_lives_matter",
        "daca and dapa": "daca_dapa",
        "lgbtq rights": "lgbtq_rights",
        "supreme court vacancy": "supreme_court_vacancy",
        # Add more replacements as needed
    }
    for key, value in replacements.items():
        text = text.replace(key, value)
    
    # Removing special characters and digits
    text = re.sub(r'[^a-zA-Z\s]', '', text)
    
    # Tokenizing using TreebankWordTokenizer
    words = tokenizer.tokenize(text)
    
    # Removing stopwords
    words = [word for word in words if word not in stop_words]
    
    # Lemmatizing
    words = [lemmatizer.lemmatize(word) for word in words]
    
    # Reconstruct the cleaned text
    cleaned_text = ' '.join(words)
    return cleaned_text


# Define candidates and topics for 2016, 2020, 2024
candidates = {
    "2016": {
        "Donald Trump": [
            "Abortion", "Black Lives Matter", "Branches of government", "Civil liberties", "Climate change",
            "Constitution", "Crime and justice", "DACA and DAPA", "Education", "Energy and environmental policy",
            "Epidemic control", "Federal assistance programs", "LGBTQ rights", "Gun control", "Healthcare",
            "Immigration", "Infrastructure", "Marriage equality", "Puerto Rico", "Stop and frisk", 
            "Transgender restroom access", "Zika virus", "Taxes", "Banking policy", "Government regulations",
            "International trade", "Budgets", "Labor and employment", "Foreign affairs", "Gay rights"
        ],
        "Hillary Clinton": [
            "Tenure as U.S. senator", "Tenure as secretary of state", "Paid speeches", 
            "backed a resolution to authorize military force in Iraq", "Abortion", "Black Lives Matter",
            "Branches of government", "Civil liberties", "Constitution", "Crime and justice", "DACA and DAPA",
            "Education", "Energy and environmental policy", "Epidemic control", "Federal assistance programs",
            "LGBTQ rights", "Gun control", "Healthcare", "Immigration", "Infrastructure", "Marriage equality",
            "Puerto Rico", "Stop and frisk", "Transgender restroom access", "Zika virus", 
            "Banking policy", "Government regulations", "International trade", "Budgets", "Labor and employment",
            "Foreign affairs", "Gay rights"
        ],
        "Gary Johnson": [
            "Abortion", "Branches of government", "Civil liberties", "Constitution", "Education", 
            "Energy and environmental policy", "Federal assistance programs", "LGBTQ rights", "Gun control",
            "Healthcare", "Immigration", "Infrastructure", "Stop and frisk", "Supreme Court vacancy",
            "Iran nuclear deal", "ISIS and terrorism", "Military and veterans", "National security"      
        ],
        "Jill Stein": [
            "Abortion", "Black Lives Matter", "Civil liberties", "Constitution", "DACA and DAPA", "Education",
            "Energy and environmental policy", "Epidemic control", "Federal assistance programs", "LGBTQ rights",
            "Gun control", "Healthcare", "Immigration", "Infrastructure", "Marijuana", "Puerto Rico", "Stop and frisk",
            "Supreme Court vacancy"      
        ]
    },
    
    "2020": {
        "Donald Trump": [
            "Immigration", "Healthcare", "Energy and environmental issues", "Trade", "Economy",
            "Education", "Gun regulation", "Criminal justice", "Foreign policy", "Impeachment",
            "Labor", "Abortion", "Supreme Court vacancy", "Other policy positions"
        ],
        "Joe Biden": [
            "Immigration", "Healthcare", "Energy and environmental issues", "Trade", "Economy",
            "Education", "Gun regulation", "Criminal justice", "Foreign policy", "Impeachment",
            "Labor", "Abortion", "Supreme Court vacancy", "Other policy positions"
        ],
        "Howie Hawkins": [
            "Immigration", "Healthcare", "Energy and environmental issues", "Trade", "Economy",
            "Education", "Gun regulation", "Criminal justice", "Foreign policy", "Impeachment",
            "Abortion", "Supreme Court vacancy", "Other policy positions"
        ],
        "Jo Jorgensen": [
            "Immigration", "Healthcare", "Energy and environmental issues", "Trade", "Economy",
            "Education", "Gun regulation", "Criminal justice", "Foreign policy", "Impeachment",
            "Abortion", "Supreme Court vacancy", "Other policy positions"
        ]
    },
    
    "2024": {
        "Kamala Harris": [
            "Immigration", "Healthcare", "Energy and environmental issues", "Economy",
            "Education", "Gun regulation", "Criminal justice", "Foreign policy", "Abortion",
            "Election policy", "Sex and gender issues", "Opioids and drug issues", "Veterans",
            "Other policy positions"
        ],
        "Donald Trump": [
            "Immigration", "Healthcare", "Energy and environmental issues", "Trade", "Economy",
            "Education", "Gun regulation", "Criminal justice", "Foreign policy", "Abortion",
            "Administrative State", "Coronavirus response", "Election policy",
            "Environmental, social, and corporate governance (ESG)", "Government ethics",
            "Sex and gender issues", "Infrastructure", "Opioids and drug issues", "Veterans",
            "Other policy positions"
        ],
        "Jill Stein": [
            "Immigration", "Healthcare", "Energy and environmental issues", "Trade", "Economy",
            "Education", "Gun regulation", "Criminal justice", "Foreign policy", "Abortion",
            "Election policy", "Government ethics", "Sex and gender issues", "Infrastructure",
            "Opioids and drug issues", "Veterans", "Other policy positions"
        ],
        "Chase Oliver": [
            "Immigration", "Healthcare", "Energy and environmental issues", "Trade", "Economy",
            "Education", "Gun regulation", "Criminal justice", "Foreign policy", "Abortion",
            "Coronavirus response", "Government ethics", "Other policy positions"
        ]
    }
}


# 
def extract_candidate_data(doc, year):
    candidate_data = {candidate: defaultdict(str) for candidate in candidates[year].keys()}
    current_candidate = None
    current_topic = None

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        
        if not text:
            continue  # Skip empty lines
        
        # Check if the text matches any candidate name (case-insensitive)
        candidate_match = next((cand for cand in candidates[year] if cand.lower() == text.lower()), None)
        if candidate_match:
            current_candidate = candidate_match
            current_topic = None  # Reset topic when candidate changes
            continue  # Move to next paragraph
        
        # Check if the text matches any topic within the current candidate's topics
        if current_candidate:
            topics = candidates[year][current_candidate]
            topic_match = next((topic for topic in topics if topic.lower() == text.lower()), None)
            if topic_match:
                current_topic = topic_match
                continue  # Move to next paragraph
        
        # If we have a candidate and topic, preprocess and append the text to the respective section
        if current_candidate and current_topic:
            processed_text = preprocess_text(text)
            candidate_data[current_candidate][current_topic] += processed_text + " "
    
    return candidate_data



# Loading Documents
def load_document(path):
    try:
        return Document(path)
    except Exception as e:
        print(f"Error loading document at {path}: {e}")
        return None

doc_2016 = load_document(r"C:\Users\az2088\OneDrive - UNC-Wilmington\Documents\DSCLLM\2016_presidential_candidate_stances.docx")
doc_2020 = load_document(r"C:\Users\az2088\OneDrive - UNC-Wilmington\Documents\DSCLLM\2020_presidential_candidate_stances.docx")
doc_2024 = load_document(r"C:\Users\az2088\OneDrive - UNC-Wilmington\Documents\DSCLLM\2024_presidential_candidate_stances.docx")

# Extracting Candidate Data
candidate_data_2016 = extract_candidate_data(doc_2016, '2016') if doc_2016 else {}
candidate_data_2020 = extract_candidate_data(doc_2020, '2020') if doc_2020 else {}
candidate_data_2024 = extract_candidate_data(doc_2024, '2024') if doc_2024 else {}


# Saving Processed Data
def save_data(data, filename):
    with open(filename, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=4)

save_data(candidate_data_2016, 'candidate_data_2016.json')
save_data(candidate_data_2020, 'candidate_data_2020.json')
save_data(candidate_data_2024, 'candidate_data_2024.json')



# Loading the documents
# Base directory and file paths
base_dir = r"C:\Users\az2088\OneDrive - UNC-Wilmington\Documents\DSCLLM"
doc_paths = {
    "2016": os.path.join(base_dir, "2016_presidential_candidate_stances.docx"),
    "2020": os.path.join(base_dir, "2020_presidential_candidate_stances.docx"),
    "2024": os.path.join(base_dir, "2024_presidential_candidate_stances.docx")
}

# Verify file paths
for year, path in doc_paths.items():
    print(f"{year} file path: {path}")


# Extract Data from .docx files
def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        # Join non-empty paragraphs to form the document content
        return "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return ""

# Extract text for each year
extracted_data = {}
for year, path in doc_paths.items():
    extracted_data[year] = extract_text_from_docx(path)
    print(f"\n{year} data:\n{extracted_data[year][:500]}...")         # Preview first 500 characters


# Define Standard Candidates names for each year
standard_candidates = {
    "2016": ["Donald Trump", "Hillary Clinton", "Gary Johnson", "Jill Stein"],
    "2020": ["Donald Trump", "Joe Biden", "Howie Hawkins", "Jo Jorgensen"],
    "2024": ["Kamala Harris", "Donald Trump", "Jill Stein", "Chase Oliver"]
}


# Parse Candidate Data
logging.basicConfig(level=logging.WARNING)

def parse_candidates_from_text(year, text):
    candidates_data = {}
    current_candidate = None

    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue  # Skip empty lines

        # Detect candidate names based on standard_candidates
        if line in standard_candidates[year]:
            current_candidate = line
            if current_candidate in candidates_data:
                logging.warning(f"Duplicate entry detected for {current_candidate}")
            candidates_data.setdefault(current_candidate, [])
        elif current_candidate:  # Add topics to the current candidate
            candidates_data[current_candidate].append(line)
        else:
            logging.warning(f"Unmatched line: {line}")

    return candidates_data


# Parse data for each year
candidates = {}
for year, text in extracted_data.items():
    candidates[year] = parse_candidates_from_text(year, text)

# Display parsed data for 2024
print("\nParsed 2024 Data:")
for candidate, topics in candidates["2024"].items():
    print(f"{candidate}: {len(topics)} topics")
    print(f"  Example topics: {topics[:5]}")


# Note: Implement Standardized Topics: 
# - From above, Kamala Harris: 29 topics, Donald Trump: 40 topics, Jill Stein: 34 topics, and Chase Oliver: 26 topics
# - So, we need to Extract the most common themes across all candidates.


# Standard Set of Topics that all candidates have 
standard_topics = [
    "Immigration", "Healthcare", "Energy and environmental issues", "Economy",
    "Education", "Gun regulation", "Criminal justice", "Foreign policy", "Abortion", "Other policy positions"
]

# Display Standard Topics
# Map candidate statements to standardized topics
def map_topics_to_standard(candidate_data, standard_topics):
    topic_mapping = {topic: [] for topic in standard_topics}

    for statement in candidate_data:
        matched = False
        for topic in standard_topics:
            if topic.lower() in statement.lower():
                topic_mapping[topic].append(statement)
                matched = True
                break
        if not matched:
            topic_mapping.setdefault("Other", []).append(statement)

    return topic_mapping



# Apply Map topics for all candidates in 2024
standardized_candidates = {}
for candidate, topics in candidates["2024"].items():
    standardized_candidates[candidate] = map_topics_to_standard(topics, standard_topics)

# Display mapped topics for one candidate
print("\nMapped Topics for Kamala Harris:")
for topic, statements in standardized_candidates["Kamala Harris"].items():
    print(f"{topic}: {len(statements)} statements")
    print(f"  Example: {statements[:1]}")  # Show an example statement



################### APP ######################

#import streamlit as st
from transformers import AutoTokenizer, AutoModelForSequenceClassification, pipeline

# Load NLI model and tokenizer
model_name = "roberta-large-mnli"
tokenizer = AutoTokenizer.from_pretrained(model_name)
model = AutoModelForSequenceClassification.from_pretrained(model_name)
nli_pipeline = pipeline("text-classification", model=model, tokenizer=tokenizer)

# Standardized topics for candidates
standard_topics = [
    "Immigration", "Healthcare", "Energy and environmental issues", "Economy",
    "Education", "Gun regulation", "Criminal justice", "Foreign policy", 
    "Abortion", "Other policy positions"
]

# Example candidate data (to be replaced with actual data)
standardized_candidates = {
    "Kamala Harris": {topic: [] for topic in standard_topics},
    "Donald Trump": {topic: [] for topic in standard_topics},
    "Jill Stein": {topic: [] for topic in standard_topics},
    "Chase Oliver": {topic: [] for topic in standard_topics},
}

# Choices for user input
choices = {
    "Immigration": [
        "I support open borders and believe all immigrants should be granted a path to citizenship.",
        "I advocate for strict border security, including a complete wall and deportation of undocumented immigrants.",
        "We need both strong border security and a fair system for legal immigration.",
        "No opinion on this topic."
    ],
    "Healthcare": [
        "I believe in a universal healthcare system funded by the government.",
        "Healthcare should be fully privatized with minimal government involvement.",
        "A hybrid system combining private insurance with government-subsidized programs is the best approach.",
        "No opinion on this topic."
    ],
    "Energy and environmental issues": [
    "We must transition immediately to 100% renewable energy sources and ban fossil fuels to combat climate change.",
    "Fossil fuels are essential to our economy, and environmental regulations should be minimized to promote business growth.",
    "We need to balance renewable energy investments with responsible use of fossil fuels while minimizing environmental impact.",
    "No opinion on this topic."
    ],
"Economy": [
    "We should heavily tax the wealthy and corporations to fund public programs like universal basic income and social services.",
    "Lower taxes for everyone, especially businesses, to promote economic growth and reduce government intervention in the market.",
    "A balanced approach to taxation that supports social programs while encouraging business growth is key to economic stability.",
    "No opinion on this topic." 
],
    
"Education": [
   "Public college should be free, and student loan debt should be canceled to make education accessible to everyone.",
    "Education funding should prioritize charter schools and private options, reducing government control over education.",
    "We should improve public education while providing support for alternative education methods like charter and private schools.",
    "No opinion on this topic."
],
    
"Gun regulation": [
    "We need a complete ban on civilian ownership of firearms, with only law enforcement and military allowed to possess guns.",
    "The Second Amendment is non-negotiable, and there should be no restrictions on gun ownership.",
    "We should protect Second Amendment rights while implementing common-sense measures like universal background checks.",
    "No opinion on this topic."
],
   
 "Criminal justice": [
    "Defund the police and redirect funds to community programs and rehabilitation efforts.",
    "Increase funding for law enforcement and implement stricter penalties for crimes to maintain law and order.",
    "We need to reform policing practices while ensuring public safety through adequate law enforcement funding.",
    "No opinion on this topic."
],
        
"Foreign policy": [
    "The U.S. should reduce military spending and focus on diplomacy and international collaboration.",
    "The U.S. must prioritize military dominance and act unilaterally to protect national interests.",
    "We need a strong military while emphasizing diplomacy and strategic alliances.",
    "No opinion on this topic."
],

"Abortion": [
    "Abortion is a fundamental right, and access should be unrestricted and fully funded by the government.",
    "Abortion should be banned in all cases, as life begins at conception.",
    "Abortion should be safe, legal, and rare, with certain restrictions like parental notification and late-term bans.",
    "No opinion on this topic."
],
    
# "Election policy": [
#     "We should make Election Day a national holiday and expand access to mail-in voting and same-day voter registration.",
#     "We must implement stricter voter ID laws and limit mail-in voting to prevent election fraud.",
#     "Election integrity requires a mix of access, like early voting, and secure processes like voter ID verification.",
#     "No opinion on this topic."

# ],
    
# "Sex and gender issues": [ 
#     "We should protect transgender rights and allow people of all ages to access gender-affirming care.",
#     "Laws should only recognize biological sex, not gender identity.",
#     "We should respect people's choices while balancing laws to protect everyone fairly.",
#     "I have no opinion on this topic."

# ],
    
# "Opioids and drug issues": [
#    "We should treat drug addiction as a health issue, not a crime, and provide free treatment and safe use facilities.",
#     "We need tougher penalties for drug users and dealers to stop the spread of opioids.",
#     "We should combine strict enforcement against dealers with better addiction treatment programs.",
#     "I have no opinion on this topic."
# ],
    
"Other policy positions": [
   "We should focus on equality and fairness for everyone, even if it means big changes to the system.",
    "We should prioritize tradition and protect the way things have always worked.",
    "We should make changes where needed but keep what already works well.",
    "I have no opinion on this topic."
]
}


def match_user_to_candidates(user_input, standardized_candidates):
    results = {}
    for candidate, topics in standardized_candidates.items():
        results[candidate] = {}
        for topic, statements in topics.items():
            if topic in user_input:
                user_stance = user_input[topic]
                combined_statements = " ".join(statements)  # Combine candidate stances
                if combined_statements:  # Avoid empty premise
                    result = nli_pipeline(
                        f"Premise: {combined_statements} Hypothesis: {user_stance}",
                        truncation=True
                    )[0]
                    score = result["score"]
                    if score > 0.50:
                        label = "ENTAILMENT"
                    elif score == 0.50:
                        label = "NEUTRAL"
                    else:
                        label = "CONTRADICTION"
                    results[candidate][topic] = {"label": label, "score": score}
    return results


def aggregate_scores(match_results):
    overall_scores = []
    for candidate, topics in match_results.items():
        total_score = sum(result["score"] for result in topics.values())
        average_score = total_score / len(topics) if topics else 0
        overall_scores.append((candidate, average_score))
    overall_scores.sort(key=lambda x: x[1], reverse=True)
    return overall_scores



################### Gradio app ######################


# Run Gradio app: app_voter_match.py

import gradio as gr
from transformers import AutoTokenizer, AutoModelForSequenceClassification, pipeline

#gr.title("🗳️ Political Candidate Match App")
#gr.header("Welcome to the 2024 Political Candidate Match App. Find out which candidate aligns most with your views")



# Load NLI model and tokenizer
model_name = "roberta-large-mnli"
tokenizer = AutoTokenizer.from_pretrained(model_name)
model = AutoModelForSequenceClassification.from_pretrained(model_name)
nli_pipeline = pipeline("text-classification", model=model, tokenizer=tokenizer)

# Preprocessed candidate data (add real data here or load from JSON)
# This is from "candidate_data_2024.json"

standardized_candidates = {
    "Kamala Harris": {
        "Immigration": "harris campaign website said vice president harris governor walz believe tough smart solution secure border keep community safe reform broken immigration system attorney general california vice president harris went international drug gang human trafficker cartel smuggled gun drug human being across usmexico border vice president supported bipartisan border security bill strongest reform decade legislation would deployed detection technology intercept fentanyl drug added border security agent protect border donald trump killed border deal political gain president biden took action ownand border crossing lowest level year administration seizing record amount fentanyl secured funding significant increase border agent ten year president bring back bipartisan border security bill sign law time know immigration system broken need comprehensive reform includes strong border security earned pathway citizenshipsource ",
        "Healthcare": "harris campaign website said vice president harris make affordable health care right privilege expanding strengthening affordable care act making permanent bidenharris tax credit enhancement lowering health care premium average year million american shell build bidenharris administration success bringing cost lifesaving prescription drug medicare beneficiary extending cap insulin cap outofpocket spending senior american tiebreaking vote inflation reduction act gave medicare power go toe toe big pharma negotiate lower drug price president shell accelerate negotiation cover drug lower price american vice president also announced medical debt removed credit report helped cancel billion medical debt million american president shell work state cancel medical debt even american vice president harris led administration effort combat maternal mortality woman nationwide dying childbirth higher rate developed nation vice president called state extend medicaid postpartum coverage two month twelve today state soup three near administration startsource ",
        "Energy and environmental issues": "harris campaign website said attorney general kamala harris ten million settlement big oil held polluter accountable vice president cast tiebreaking vote pas inflation reduction act largest investment climate action history historic work lowering household energy cost creating hundred thousand highquality clean energy job building thriving clean energy economy ensuring america energy security independence record energy production president unite american tackle climate crisis build historic work advance environmental justice protects public land public health increase resilience climate disaster lower household energy cost creates million new job continues hold polluter accountable secure clean air water vice president said international climate conference cop know meeting global challenge require global cooperation committed continuing building upon united state international climate leadership governor walz always fight freedom breathe clean air drink clean water live free pollution fuel climate crisissource ",
        "Economy": "harris campaign website said vice president harris grew middle class home daughter working mom belief middle class strong america strong thats president kamala harris create opportunity economy everyone chance compete chance succeedwhether live rural area small town big city vice president kamala harris made clear building middle class defining goal presidency thats make top priority bring cost increase economic security american president fight cut tax million working middle class american lowering cost everyday need like health care housing grocery bring together organized labor worker small business owner entrepreneur american company create good paying job grow economy ensure america continues lead worldsource ",
        "Education": "harris campaign website said vice president harris fight ensure parent afford highquality child care preschool child strengthen public education training pathway middle class shell continue working end unreasonable burden student loan debt fight make higher education affordable college ticket middle class date vice president harris helped deliver largest investment public education american history provide nearly billion student debt relief almost five million borrower deliver record investment hbcus tribal college hispanicserving institution minorityserving institution helped student afford college increasing maximum pell grant award largest increase decadeand invested community college implemented policy led one million registered apprentice hired even scale program create good career pathway noncollege graduatessource ",
        "Gun regulation": "harris campaign website said president biden vice president harris encouraged bipartisan cooperation pas first major gun safety law nearly year included record funding hire train mental health professional school head firstever white house office gun violence prevention spearheaded policy expand background check close gun show loophole president wont stop fighting american freedom live safe gun violence school community place worship shell ban assault weapon highcapacity magazine require universal background check support red flag law keep gun hand dangerous people also continue invest funding law enforcement including hiring training officer people support build upon proven gun violence prevention program helped reduce violent crime throughout countrysource ",
        "Criminal justice": "harris campaign website said prosecutor vice president harris fought violent crime getting illegal gun violent criminal california street time district attorney raised conviction rate violent offendersincluding gang member gun felon domestic abuser attorney general vice president harris built record removing illegal gun street california prosecuting toughest transnational criminal organization world white house vice president harris helped deliver largest investment public safety ever investing billion supporting local law enforcement community safety program across city town county president also continue invest funding law enforcement including hiring training officer people support build upon proven gun violence prevention program helped reduce violent crime throughout countrysource ",
        "Foreign policy": "harris campaign website said vice president harris never waver defense america security ideal vice president confronted threat security negotiated foreign leader strengthened alliance engaged brave troop overseas commander chief ensure united state military remains strongest lethal fighting force world unleash power american innovation win competition st century strengthen abdicate global leadership vice president harris fulfill sacred obligation care troop family always honor service sacrificesource ",
        "Abortion": "harris campaign website said vice president harris governor walz trust woman make decision body government tell since roe v wade overturned vice president harris driven administration strategy defend reproductive freedom safeguard privacy patient provider governor tim walz led minnesota become first state pas law protecting woman right choose following overturning roe vice president harris traveled america heard story woman hurt trump abortion ban story couple trying grow family cut middle ivf treatment story woman miscarrying parking lot developing sepsis losing ability ever child doctor afraid may go jail caring patient president never allow national abortion ban become law congress pass bill restore reproductive freedom nationwide sign itsource ",
        "Election policy": "harris campaign website said vice president harris governor walz believe many fundamental freedom stake election fight ensure american opportunity participate democracy passing john lewis voting right freedom vote actslaws enshrine voting right protection expand votebymail early voting moresource ",
        "Sex and gender issues": "harris campaign website said president shell always defend freedom love love openly pride officiated nation first samesex marriage attorney general refused defend california antimarriage equality statewide referendum president shell fight pas equality act enshrine antidiscrimination protection lgbtqi american health care housing education lawsource ",
        "Opioids and drug issues": "harris campaign website said vice president harris committed ending opioid epidemic tackling scourge fentanyl shes seen devastating impact fentanyl family closeshe met mourned lost loved one fentanyl overdoses attorney general prosecuted drug trafficker seizing kilo cocaine pound methamphetamine white house helped direct billion disrupt flow illicit drug delivered billion dollar investment state fund lifesaving program bidenharris administration fda made overdosereversal drug naloxone available overthecounter past year number overdose death united state declined first time five year president sign bipartisan border bill fund detection technology intercept even illicit drug shell keep fighting end opioid epidemicsource ",
        "Veterans": "harris campaign website said vice president harris governor walz believe sacred obligation care nation service member veteran family caregiver survivor american represent bravest among u put life line defend promise america stand value protect fundamental freedom vice president harris president biden delivered significant expansion benefit service veteran exposed burn pit agent orange toxic substance year son army veteran served command sergeant major governor walz ranking member house veteran affair committee passed legislation help stem veteran suicide fight end veteran homelessness investing mental health suicide prevention effort eliminating barrier employment expanding economic opportunity military veteran family harriswalz administration continue ensure service member veteran family receive benefit earnedsource additional reading ",
        "Other policy positions": "click following link read policy position presidential candidate "
    },
    "Donald Trump": {
        "Immigration": "trump campaign website said president trump shut bidens border disaster end catchandrelease restore remain mexico eliminate asylum fraud cooperative state president trump deputize national guard local law enforcement assist rapidly removing illegal alien gang member criminal also deliver meritbased immigration system protects american labor promotes american valuessource ",
        "Healthcare": "trump campaign website said president donald j trump empowered american patient greatly expanding healthcare choice transparency affordability increased competition health insurance market eliminated obamacare individual mandate signed right try give terminally ill patient access lifesaving cure president trump lowered drug price first time year finalized favored nation rule ensure pharmaceutical company offer discount united state nationssource ",
        "Energy and environmental issues": "trump campaign website said joe biden reversed trump energy revolution enriching foreign adversary abroad president trump unleash production domestic energy resource reduce soaring price gasoline diesel natural gas promote energy security friend around world eliminate socialist green new deal ensure united state never mercy foreign supplier energysource ",
        "Trade": "trump campaign website said heart vision sweeping proamerican overhaul tax trade policy move biden system punishes domestic producer reward outsourcers system reward domestic production tax foreign company export american job rewarded rewarded greatly country benefit achieve goal phase system universal baseline tariff foreign product top higher tariff increase incrementally depending much individual foreign country devalue currencysource ",
        "Economy": "trump campaign website said president donald j trump passed recordsetting tax relief middle class doubled child tax credit slashed jobkilling regulation administration ever done real wage quickly increased result median household income reached highest level history country poverty reached record low president trump created nearly opportunity zone revitalize neglected community president trump produced booming economic recovery record low unemployment african american hispanic american asian american woman joe biden destroyer america job continues fuel runaway inflation reckless big government spending president trump vision america economic revival lower tax bigger paycheck job american workerssource ",
        "Education": "trump campaign website said president trump belief owe child great school lead great job lead even greater country living right end president trump work ensure top priority every school prepare student job connection totally refocusing school succeeding world work president trump pledge close department education washington dc send education work need back statessource ",
        "Gun regulation": "trump campaign website said also always defend second amendment right keep bear armssource ",
        "Criminal justice": "trump campaign website said higher priority quickly restoring law order public safety america president trump stand hero law enforcement joe biden radical left politician defunded defamed dismantled police force across america murder spiked alltime high democratrun city radical prosecutor district attorney given free rein violent criminal threaten citizen street oncegreat city controlled gang cartel plagued mentally ill drugaddicted homelesssource ",
        "Foreign policy": "trump campaign website said president donald j trump replaced failed policy neverending war regime change nationbuilding bold vision pursue peace strength joe biden undermined military readiness surrendered strength taliban president trump defend america threat protect america danger keep america unnecessary foreign war also get bidens radical left ideology military rehire every patriot unjustly fired protect people threat nuclear weapon hypersonic missile president trump also build stateoftheart nextgeneration missile defense shieldsource ",
        "Abortion": "trump said statement proudly person responsible ending something legal scholar side wanted fact demanded ended roe v wade wanted ended view abortion everybody wanted legal standpoint state determine vote legislation perhaps whatever decide must law land case law state many state different many different number week conservative law others thats end day peoplesource ",
        "Administrative State": "trump campaign website listed following policy day one reissue executive order restoring president authority fire rogue bureaucrat overhaul federal department agency firing corrupt actor national security intelligence apparatus launch major crackdown government leaker collude medium create false narrative pressing criminal charge appropriate make every inspector general office independent department oversee become protector deep state trump campaign website also listed continue trump administration effort move part federal bureaucracy outside washington swamp like president trump moved bureau land management colorado government position could moved washington ban federal bureaucrat taking job company deal regulate big pharmasource ",
        "Coronavirus response": "trump campaign website said save life china virus president trump organized production world largest supply ventilator development treatment vaccine stop covid mandate restore medical freedomsource ",
        "Election policy": "trump campaign website said president donald j trump committed honesty election integrity republic reform election law verify identity eligibility voter ensure faith confidence future election pas bold range critical election integrity measure include banning unsecure drop box ballot harvesting state local official permitted make illegal unconstitutional change election procedure without required approval state legislature importantly must ban private money pouring local election officessource ",
        "Environmental, social, and corporate governance (ESG)": "trump campaign website said entire esg scheme designed funnel retirement money maniac radical left rule issued leadership first esg ban anywhere world im delighted republican congress across country waking threat following leadsource ",
        "Government ethics": "trump campaign website said president donald j trump committed dismantling deep state restoring government people administration president trump conduct toptobottom overhaul federal bureaucracy clean rot corruption washington dc president trump push constitutional amendment impose term limit member congress permanent ban taxpayer funding campaign lifetime ban lobbying former member congress cabinet member ban member congress trading stock insider informationsource ",
        "Sex and gender issues": "trump campaign website said ask congress pas bill establishing gender recognized united state government male femaleand assigned birth bill also make clear title nine prohibits men participating womens sport protect right parent forced allow minor child assume gender new identity without parent consent identity new without parental consentsource ",
        "Infrastructure": "trump campaign website said word well actually build new city country freedom city reopen frontier reignite american imagination give hundred thousand young people people hardworking family new shot home ownership fact american dream challenge governor state join great modernization beautification campaigngetting rid ugly building refurbishing park public space making city town livable ensuring pristine environment building towering monument true american heroessource ",
        "Opioids and drug issues": "trump campaign website said president donald j trump marshalled full power government stop deadly drug opioids fentanyl coming country result drug overdose death declined nationwide first time nearly year joe biden allowed drug cartel wage war america steal innocent life ravage community president trump take drug cartel took isissource ",
        "Veterans": "trump campaign website said president donald j trump passed largest reform department veteran affair generation including va accountability va choice fired federal worker failed give wounded warrior quality timely care richly deserve secured record funding mental health service expanded access telehealth suicide prevention resource secure blessing freedom risked life defend president trump decreased veteran homelessness increased educational benefit achieved recordlow veteran unemploymentsource ",
        "Other policy positions": "click following link read policy position presidential candidate "
    },
    "Jill Stein": {
        "Immigration": "stein campaign website said time completely overhaul broken abusive immigration system well unjust policy driving people leave home need comprehensive immigration policy properly funded institution ensure timely ethical transparent dignified path citizenship immigrant asylum seeker refugee border policy move away detention enforcement response toward humane effective asylum processing includes full support funding coordinated civil society response including social legal service provider instead jailing migrant asylum seeker create noncustodial humanitarian reception center border migrant processed rapidly screen significant criminal record processed migrant paper begin work immediately making invaluable resource communitiessource ",
        "Healthcare": "stein campaign website said healthcare system crisis united state spends healthcare highincome country worse health outcome including lowest life expectancy birth highest rate people multiple chronic disease million people uninsured u many insured still cant afford healthcare due huge outofpocket cost researcher estimate lack adequate healthcare led excess death covid wall street party funded insurance industry pharmaceutical industry big healthcare profiteer perpetuate failed system put profit people healthcare human right need universal healthcare system equitable comprehensive free point service accessible every single person ussource ",
        "Energy and environmental issues": "stein campaign website said today face worsening global climate crisis threatens future human life earth hottest year record hottest year record occurred past decade accelerating global heating wreaking havoc ecosystem leading bigger deadlier fire flood megastorms drought heat wave environmental climate injustice disproportionately harming black brown lowincome indigenous community across country around world yet despite existential emergency last year seen massive expansion fossil fuel planning infrastructure production breaking record u oil production wall street party funded corporate polluter driving u climate cliff human right livable planet stable climate healthy food clean air water living soil need real green new deal transition rapidly economic system destroying home sustainable society built around human need protecting life earth need act child future generation survive thrivesource ",
        "Trade": "stein campaign website listed following policy replace corporate trade agreement global fair trade agreementssource ",
        "Economy": "stein campaign website said need economy work working people wealthy powerful reverse surging inequality insecurity need economic bill right establishing right livingwage job guaranteed livable income housing healthcare childcare lifelong education secure retirement utility healthy food clean water u guaranteed basic security good life reach highest potentialsource ",
        "Education": "stein campaign website listed following policy guarantee lifelong free public education institution learning including trade school prek college graduate school abolish student debt million encumbered american increase equalize public school funding end privatization public schoolssource ",
        "Gun regulation": "stein campaign website listed following policy end epidemic gun violence commonsense gun safety law ban sale assault rifle establish buyback program establish mandatory waiting period background check firearm purchase pas red flag law individual pose danger others create standardized digital record gun registration sale close gun show loophole require firearm owner highquality gun safe store firearm require firearm owner purchase liability insurance less hold adult firearm owner criminally liable minor child accessing firearm using commission crime accidental injury deathsource ",
        "Criminal justice": "stein campaign website said call united state land free highest incarceration rate country world million people federal state local prison jail instead addressing root cause inequality injustice today system policing prison criminal justice designed wall street party wealthy elite backer enforce socioeconomic hierarchy systemically racist classist need end mass incarceration police brutality systemic injustice jill stein administration guarantee human right restorative criminal justice system treat every one million people federal state local prison jail respect dignity compassion primary goal reengaging family communitiessource ",
        "Foreign policy": "stein campaign website said bipartisan endless war machine enriches military contractor lobbyist politician fuel devastation around world impoverishes people pentagon budget consumes half discretionary federal budget real u military spending trillion dollar per year militaryindustrial complex aided accomplice war party medium intelligence agency beyond become global empire profoundly destructive around world home everyone human right live peace dignity free violence oppression must end endless war create new foreign policy based diplomacy international law human right lead way new era peace cooperationsource ",
        "Abortion": "stein campaign website listed following policy codify roe v wade ensure full reproductive right bodily autonomy woman repeal hyde amendmentsource ",
        "Election policy": "stein campaign website listed following policy replace exclusionary twocorporateparty system inclusive multiparty democracy rankedchoice voting proportional representation implement rankedchoice voting election nationwide implement proportional representation legislative election institute full public financing election get corrupting influence private money politics put people back abolish electoral college elect president via national popular vote using rankedchoice voting support modern voting right act including nonpartisan redistricting commission sameday voter registration nationwide restore preclearance provision voting right act ensure constitutional right vote restore voting right felon pas automatic voter registration avr nationwide make election day federal holiday expand polling location make free votebymail option election expand polling location end discriminatory voting law purging voting roll repeal shelby county v holder allow supervision incarcerated vote election counted district resided incarceration eliminate gerrymandering enacting proportional representation repeal discriminatory antidemocratic ballot access restriction designed establishment party suppress competition expand initiative referendum recall power every state nationally safeguard election integrity handcounted paper ballot routine postelection audit lower voting age source ",
        "Government ethics": "stein campaign website said democracy life support belief political system historic low number american feel neither establishment party represents record high researcher found today u government oligarchy policy determined people demand corporate elite wall street party systematically concentrated power hand wealthy donor locking people rightful place decisionmaking table full meaningful participation democracy human right need revive democracy full spectrum reform empower people including real choice ballot without freedom choice election democracysource ",
        "Sex and gender issues": "stein campaign website said jill stein administration guarantee human right slgbtqia community violent attack transgender woman particularly black transgender woman color recent stripping state protection lack legal protection negative court ruling longstanding historical inequity continue end acceptance violent culture devalues humanity slgbtqia sibling fight liberation slgbtqia people around worldsource ",
        "Infrastructure": "stein campaign website listed following policy guarantee affordable efficient utility transition utility public notforprofit ownership free highspeed internet across u rural broadband via fiber opticssource ",
        "Opioids and drug issues": "stein campaign website listed following policy begin process decriminalizing personal possession hard drug treat drug misuse health problem criminal problemsource ",
        "Veterans": "stein campaign website listed following policy fully fund veteran program benefit including healthcare mental health housing job training transition civilian lifesource ",
        "Other policy positions": "click following link read policy position presidential candidate "
    },
    "Chase Oliver": {
        "Immigration": "oliver campaign website said immigration system tragic mess simply must break partisan logjam bring people party together reform modernize better example bad government overly complex current law regulation involving immigration president need take lead work congress radically simplify immigration system people come work become member community without relegated shadow immigrant built country help keep growing prospering future generation americanssource ",
        "Healthcare": "oliver campaign website said healthcare expensive government overregulation advocate market alternative heavily regulated employerprovided insurance direct primary care model also reform bring new drug market end practice patent evergreening keep cost many drug artificially high choice free market force see drop cost healthcare overallsource ",
        "Energy and environmental issues": "debate oliver said need clean air clean water company poison community able sue company without tort cap let citizen jury bankrupt polluter free market incentive make sure never happening againsource ",
        "Trade": "oliver campaign website said would immediately end tariff serve increase bottom line protected industry shift labor efficient industry creating net loss job raise price lowering number alternative consumer choose among tariff form embargo domestic laborer consumer raising cost business importer pas added cost onto endbuyerssource ",
        "Economy": "oliver campaign website said every day average american face many challenge overly regulated economy government lobbyist special interest continuously conspire pick winner loser trillion dollar government deficit trillion unfunded liability government recklessly crowd funding would better served allowing innovation marketplace find solution pressing issue combat constant round economic financial crisis must get government boardroom wallet allow individual decide best distribute hardearned dollarssource ",
        "Education": "oliver campaign website said federal level best policy education remove federal government involvement education support abolishing department education block granting fund back state returned taxpayer also support getting government student loan business market force lower cost higher education timesource ",
        "Gun regulation": "oliver campaign website said support right every person defend violence aggression gun control throughout history used limit ability vulnerable population defend armed people harder oppress harder attack support new gun restriction look repeal restriction exist today nominate judiciary judge committed protecting right person defend aggression district columbia v heller seminal moment history affirming right citizen arm determining restriction trigger lock unconstitutional work extend legal concept ban ban bump stockssource ",
        "Criminal justice": "oliver campaign website said long past time address problem criminal justice system court police prison need reform every area one goal mind empower people government less five percent world population host roughly quarter prison population simply unacceptable nation style land free current system feed prisonindustrial complex lead militarization police corruption prosecutor perverse incentive push longer sentence building private prison value profit rehabilitation cycle incarceration breed lowered earnings poverty lead recidivismsource ",
        "Foreign policy": "oliver campaign website said entire adult life nation war somewhere world friend loved one served come home visible invisible scar combat ive seen mourn brother sister didnt come home meet antiwar veteran every day know anyone cost war time focus foreign policy peace end policy sending drone around world instead foster international goodwill defending free trade free market nation long moniker leader free world time earn distinction insisting peace way forward end war end drone end policy constant intervention easy drop bomb much harder serve beacon peace must take difficult necessary pathsource ",
        "Abortion": "oliver campaign website said outside obvious criminal offense assault theft murder role government free society prohibit behavior even majority wish reason activity abortion province individual choice bodily autonomy recognize complexity issue good faith position side abortion debate said work encourage state decriminalize abortion decision rest hand individual doctor encourage passage hyde amendment prevent federal dollar funding abortion clinic encourage state incentivize alternative abortion adoption easing burdensome regulation greatly increase cost median range adopting newborn baby placing avenue range many would otherwise willingsource ",
        "Coronavirus response": "debate oliver said im elected president im going appoint special counsel department justice investigate everything gone wrong basically punish wrongdoer find committed force fraud coercion punish people free people able make medical decision dont need state apparatus telling u trying lock u control u one greatest violation liberty life saw million business affected arbitrary lockdown didnt stop disease saw attempt osha mandate behavior individual businessowners property owner able determine frankly individual able determine themselvessource ",
        "Government ethics": "debate oliver said powerful interest always controlled government beginning time government immoral strike government much possiblesource ",
        "Other policy positions": "click following link read policy position presidential candidate "
    }
}


# Define topics and user choices

choices = {
    "Immigration": [
    "I support open borders and believe all immigrants, regardless of status, should be granted a path to citizenship and access to social services.",
    "I advocate for strict border security, including a complete wall and deportation of undocumented immigrants without exception.",
    "We need both strong border security and a fair system for legal immigration, including pathways to citizenship for those already contributing to society.",
    "No opinion on this topic."
 ], 
"Healthcare": [
    "I believe in a universal healthcare system funded by the government where everyone has free access to medical services.",
    "Healthcare should be fully privatized, with minimal government involvement, allowing market competition to determine cost and access.",
    "A hybrid system combining private insurance options with government-subsidized programs for those in need is the best approach.",
    "No opinion on this topic."
],
    
"Energy and environmental issues": [
    "We must transition immediately to 100% renewable energy sources and ban fossil fuels to combat climate change.",
    "Fossil fuels are essential to our economy, and environmental regulations should be minimized to promote business growth.",
    "We need to balance renewable energy investments with responsible use of fossil fuels while minimizing environmental impact.",
    "No opinion on this topic."
], 
"Economy": [
    "We should heavily tax the wealthy and corporations to fund public programs like universal basic income and social services.",
    "Lower taxes for everyone, especially businesses, to promote economic growth and reduce government intervention in the market.",
    "A balanced approach to taxation that supports social programs while encouraging business growth is key to economic stability.",
    "No opinion on this topic." 
],
    
"Education": [
   "Public college should be free, and student loan debt should be canceled to make education accessible to everyone.",
    "Education funding should prioritize charter schools and private options, reducing government control over education.",
    "We should improve public education while providing support for alternative education methods like charter and private schools.",
    "No opinion on this topic."
],
    
"Gun regulation": [
    "We need a complete ban on civilian ownership of firearms, with only law enforcement and military allowed to possess guns.",
    "The Second Amendment is non-negotiable, and there should be no restrictions on gun ownership.",
    "We should protect Second Amendment rights while implementing common-sense measures like universal background checks.",
    "No opinion on this topic."
],
   

 "Criminal justice": [
    "Defund the police and redirect funds to community programs and rehabilitation efforts.",
    "Increase funding for law enforcement and implement stricter penalties for crimes to maintain law and order.",
    "We need to reform policing practices while ensuring public safety through adequate law enforcement funding.",
    "No opinion on this topic."
],
        
"Foreign policy": [
    "The U.S. should reduce military spending and focus on diplomacy and international collaboration.",
    "The U.S. must prioritize military dominance and act unilaterally to protect national interests.",
    "We need a strong military while emphasizing diplomacy and strategic alliances.",
    "No opinion on this topic."
],

"Abortion": [
    "Abortion is a fundamental right, and access should be unrestricted and fully funded by the government.",
    "Abortion should be banned in all cases, as life begins at conception.",
    "Abortion should be safe, legal, and rare, with certain restrictions like parental notification and late-term bans.",
    "No opinion on this topic."
],
    
# "Election policy": [
#     "We should make Election Day a national holiday and expand access to mail-in voting and same-day voter registration.",
#     "We must implement stricter voter ID laws and limit mail-in voting to prevent election fraud.",
#     "Election integrity requires a mix of access, like early voting, and secure processes like voter ID verification.",
#     "No opinion on this topic."

# ],
    
# "Sex and gender issues": [ 
#     "We should protect transgender rights and allow people of all ages to access gender-affirming care.",
#     "Laws should only recognize biological sex, not gender identity.",
#     "We should respect people's choices while balancing laws to protect everyone fairly.",
#     "I have no opinion on this topic."

# ],
    
# "Opioids and drug issues": [
#    "We should treat drug addiction as a health issue, not a crime, and provide free treatment and safe use facilities.",
#     "We need tougher penalties for drug users and dealers to stop the spread of opioids.",
#     "We should combine strict enforcement against dealers with better addiction treatment programs.",
#     "I have no opinion on this topic."
# ],
    
"Other policy positions": [
   "We should focus on equality and fairness for everyone, even if it means big changes to the system.",
    "We should prioritize tradition and protect the way things have always worked.",
    "We should make changes where needed but keep what already works well.",
    "I have no opinion on this topic."
]
}


candidate_urls = {
    "Kamala Harris": "https://ballotpedia.org/Kamala_Harris_presidential_campaign,_2024#Policy_positions",
    "Donald Trump": "https://ballotpedia.org/Donald_Trump_presidential_campaign,_2024#Policy_positions",
    "Jill Stein": "https://ballotpedia.org/Jill_Stein_presidential_campaign,_2024#Policy_positions",
    "Chase Oliver": "https://ballotpedia.org/Chase_Oliver_presidential_campaign,_2024#Policy_positions"
}

def match_user_to_candidates(user_input, standardized_candidates):
    results = {}
    for candidate, topics in standardized_candidates.items():
        results[candidate] = {}
        for topic, statements in topics.items():
            if topic in user_input:
                user_stance = user_input[topic]
                combined_statements = " ".join(statements)  # Combine candidate stances
                if combined_statements:  # Avoid empty premise
                    result = nli_pipeline(
                        f"Premise: {combined_statements} Hypothesis: {user_stance}",
                        truncation=True
                    )[0]
                    score = result["score"]
                    if score > 0.50:
                        label = "ENTAILMENT"
                    elif score == 0.50:
                        label = "NEUTRAL"
                    else:
                        label = "CONTRADICTION"
                    results[candidate][topic] = {"label": label, "score": score}
    return results


def aggregate_scores(match_results):
    overall_scores = []
    for candidate, topics in match_results.items():
        total_score = sum(result["score"] for result in topics.values())
        average_score = total_score / len(topics) if topics else 0
        overall_scores.append((candidate, average_score))
    overall_scores.sort(key=lambda x: x[1], reverse=True)
    return overall_scores


import time


def gradio_interface(*args):
    start_time = time.time()
    user_input = {topic: args[i] for i, topic in enumerate(choices.keys())}
    match_results = match_user_to_candidates(user_input, standardized_candidates)
    overall_scores = aggregate_scores(match_results)

    end_time = time.time()
    elapsed_time = end_time - start_time

    # Find the best matched candidate
    best_matched_candidate = overall_scores[0][0]
    best_matched_candidate_url = candidate_urls.get(best_matched_candidate, "#")

    formatted_results = "## Policy Scores of Matched Candidate\n"
    for topic, result in match_results[best_matched_candidate].items():
        formatted_results += f"- **{topic}**: {result['label']} (Score: {result['score']:.2f})\n"
    formatted_results += "\n"

    formatted_results += "### Your Overall Candidate Alignment\n"
    for candidate, score in overall_scores:
        formatted_results += f"- **{candidate}**: {score:.2f}\n"

    formatted_results += f"\n**Processing Time**: {elapsed_time:.2f} seconds\n"
    formatted_results += f"\n**To learn more about your best match Click on this URL**: [Visit {best_matched_candidate}]({best_matched_candidate_url})"

    return formatted_results



inputs = []
for topic, options in choices.items():
    inputs.append(gr.Radio(choices= options, label = f"What's your stance on {topic}?"))

outputs = gr.Markdown(label = "Match Results")

iface = gr.Interface(fn = gradio_interface, 
                     inputs = inputs, 
                     outputs = outputs, 
                     title = "🗳️ Voter's Match App ", 
                     description="Welcome to the 2024 Political Candidate Match App. Find out which candidate aligns most with your views. Please select your stance on each topic and click 'Submit' to see the results.")

iface.launch(server_port= 7861, share= True, debug= True)




# Note: The app allows users to select their stance on various topics and then compares their stances with those of the candidates.
# - The app uses a pre-trained NLI model to compare user stances with candidate statements.
# - The app then aggregates the scores for each candidate to provide an overall alignment score.
# - The user can see the alignment for each topic and the overall alignment with each candidate.
# Run : streamlit run app_candidate_match.py
